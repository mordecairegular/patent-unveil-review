"""md_to_docx should emit editable Word equations, not formula images/code text."""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))

from md_to_docx import convert_md_to_docx  # noqa: E402
from qa_docx_math import check_docx_math  # noqa: E402


def _document_xml(doc, tmp_path: Path) -> str:
    out = tmp_path / "out.docx"
    doc.save(out)
    with zipfile.ZipFile(out) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def test_code_wrapped_math_symbol_becomes_omml(tmp_path: Path) -> None:
    doc = convert_md_to_docx(
        "符号 `B_{s,t}^{tot}` 表示时段 t 内场站 s 的总储能量。",
        base_dir=tmp_path,
    )
    xml = _document_xml(doc, tmp_path)

    assert "<m:oMath" in xml
    assert "Consolas" not in xml
    assert "B_{s,t}^{tot}" not in xml


def test_block_math_with_legacy_png_comment_skips_formula_image(tmp_path: Path) -> None:
    math_dir = tmp_path / "math_figures"
    math_dir.mkdir()
    (math_dir / "eq_001.png").write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00"
        b"\x1f\x15\xc4\x89\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    md = (
        "\\[\n"
        "B_{s,t}^{tot} = B_{s,t}^{ch} + B_{s,t}^{dis} \\tag{1}\n"
        "\\]\n"
        "<!-- ![公式](math_figures/eq_001.png) -->\n"
    )

    doc = convert_md_to_docx(md, base_dir=tmp_path)
    xml = _document_xml(doc, tmp_path)

    assert "<m:oMath" in xml
    assert "word/media" not in xml
    assert "eq_001.png" not in xml


def test_fraction_math_uses_editable_omml_fraction_and_number_cell(tmp_path: Path) -> None:
    doc = convert_md_to_docx(
        "\\[\n"
        "B_{s,t}=B_{s,t-1}+\\eta_{\\mathrm{ch}}E_{s,t,\\mathrm{ch}}"
        "-\\frac{E_{s,t,\\mathrm{dis}}}{\\eta_{\\mathrm{dis}}}\\tag{1}\n"
        "\\]\n",
        base_dir=tmp_path,
    )
    xml = _document_xml(doc, tmp_path)

    assert "<m:f>" in xml
    assert "<m:sSub" in xml
    assert "frac{" not in xml
    assert "mathrm{" not in xml
    assert "(1)" in xml
    assert "<w:tbl>" in xml


def test_bare_formula_line_with_next_line_number_becomes_numbered_omml_table(tmp_path: Path) -> None:
    doc = convert_md_to_docx(
        "E(r) = V_sys / [r · ln(R_socket/R_pin)]\n"
        "(5)\n"
        "式中：E(r)为半径r处的径向电场强度，V/m；V_sys为系统直流侧电压。\n",
        base_dir=tmp_path,
    )
    xml = _document_xml(doc, tmp_path)

    assert "<m:oMath" in xml
    assert "<m:sSub" in xml
    assert "<w:tbl>" in xml
    assert "(5)" in xml
    assert "V_sys" not in xml
    assert "R_socket" not in xml
    assert "R_pin" not in xml


def test_docx_math_qa_fails_on_literal_latex_residue(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("错误公式文本 frac{E}{eta} 和 \\(p_min(s)\\)。")
    out = tmp_path / "bad.docx"
    doc.save(out)

    report = check_docx_math(out)

    assert not report.passed
    assert any(item["pattern"] == r"\(" for item in report.failed_patterns)
    assert any(item["pattern"] == "frac{" for item in report.failed_patterns)


def test_docx_math_qa_fails_on_bare_subscript_identifier(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Bad visible math residue: V_sys and R_pin should be subscripts.")
    out = tmp_path / "bad_bare_subscript.docx"
    doc.save(out)

    report = check_docx_math(out)

    assert not report.passed
    assert any(item["pattern"] == "bare_subscript_identifier" for item in report.failed_patterns)


def test_docx_delivery_qa_fails_on_unrendered_mermaid_code_block(tmp_path: Path) -> None:
    doc = convert_md_to_docx(
        "```mermaid\nflowchart TB\nA[输入] --> B[输出]\n```\n",
        base_dir=tmp_path,
    )
    out = tmp_path / "bad_mermaid.docx"
    doc.save(out)

    report = check_docx_math(out)

    assert not report.passed
    assert report.code_style_count > 0
    assert any(item["pattern"] == "mermaid_source_residue" for item in report.failed_patterns)
    assert any("Consolas" in item for item in report.structural_failures)


def test_docx_delivery_qa_fails_when_expected_media_is_missing(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("没有嵌入任何图示。")
    out = tmp_path / "missing_media.docx"
    doc.save(out)

    report = check_docx_math(out, min_media_count=1)

    assert not report.passed
    assert report.media_count == 0
    assert any("media count" in item for item in report.structural_failures)


def test_docx_math_qa_ignores_body_parenthetical_numbers_for_duplicates(tmp_path: Path) -> None:
    doc = convert_md_to_docx(
        "\\[\n"
        "E(r)=V_{sys}/r\\tag{1}\n"
        "\\]\n"
        "普通条款编号（1）和英文 parenthetical (1) 不应算作公式编号重复。\n",
        base_dir=tmp_path,
    )
    out = tmp_path / "number_scope.docx"
    doc.save(out)

    report = check_docx_math(out)

    assert report.passed
    assert report.equation_number_count == 1
    assert report.duplicate_numbers == []


def test_manifest_numbers_auto_tag_untagged_block_math(tmp_path: Path) -> None:
    doc = convert_md_to_docx(
        "\\[\n"
        "A_s=B_s\\tag{4}\n"
        "\\]\n"
        "\\[\n"
        "C_s=D_s\n"
        "\\]\n",
        base_dir=tmp_path,
        formula_numbers=["4", "5"],
    )
    out = tmp_path / "auto_numbered.docx"
    doc.save(out)
    xml = _document_xml(doc, tmp_path)

    report = check_docx_math(out)

    assert report.passed
    assert report.equation_number_count == 2
    assert report.duplicate_numbers == []
    assert "(4)" in xml
    assert "(5)" in xml


def test_docx_math_qa_passes_fixture_with_manifest(tmp_path: Path) -> None:
    fixture_dir = ROOT / "tests" / "fixtures" / "docx_math"
    md = (fixture_dir / "sample_formula_disclosure.md").read_text(encoding="utf-8")
    manifest = fixture_dir / "sample_formula_manifest.yaml"
    doc = convert_md_to_docx(md, base_dir=fixture_dir)
    out = tmp_path / "formula_fixture.docx"
    doc.save(out)

    report = check_docx_math(out, manifest_path=manifest)

    assert report.passed
    assert report.math_block_count >= 5
    assert report.equation_number_count == 5
