#!/usr/bin/env python3
"""
将 Markdown 转为 Word（.docx），按标题层级映射为 Word 内置「标题 1–9」样式，
便于交底书交付代理人或所内流程。

支持：ATX 标题 (#–######)、段落、**粗体**、行内 `代码`、无序/有序列表、
围栏代码块、简单 GFM 表格、引用块（>）、水平线（---）、行内图片 ``![](path.png)``
（在最大宽、最大高约束下**等比缩放**，竖图自动缩小宽度以整图落入版面）。

**连续多行正文**（中间无空行、且非列表/标题等）时，**每一行**输出为 Word 中**独立一段**，
以便「（1）…（2）…」等分条换行；若须在同一段内接排，请写**同一行**内或用 Markdown 空行分隔逻辑段。

定稿宜先用同目录 **`mermaid_render.py`** 将 **mermaid** 转为 PNG；**LaTeX 公式**
（``$...$`` / ``$$...$$``）由本脚本写入 Word 原生 OMML 可编辑公式。
保存 DOCX 后默认执行 ``qa_docx_math.py``，若残留 LaTeX 命令或公式结构异常则退出失败。

用法：
  python md_to_docx.py --input disclosure.md --output disclosure.docx
  python md_to_docx.py -i a.md -o b.docx --base-dir .   # 解析图片相对路径

依赖：python-docx
"""

from __future__ import annotations

import argparse
import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# 插图最大尺寸（英寸）：在常见 A4、默认边距下保证整图可见、按比例缩放（不过宽也不过高）。
_DEFAULT_IMAGE_MAX_W_IN = 5.5
_DEFAULT_IMAGE_MAX_H_IN = 8.2
# 公式图在 Word 中统一按固定高度嵌入（英寸），避免块级式随 PNG 像素被放大、行内式过小
_FORMULA_DISPLAY_MAX_H_IN = 0.17
# 兼容旧名
_FORMULA_INLINE_MAX_H_IN = _FORMULA_DISPLAY_MAX_H_IN
_FORMULA_BLOCK_MAX_W_IN = 4.0  # 仅作块级超宽时的宽度上限（通常由固定高度约束）
_FORMULA_BLOCK_MAX_H_IN = _FORMULA_DISPLAY_MAX_H_IN

_MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_HIDDEN_MD_IMAGE_COMMENT_RE = re.compile(
    r"<!--\s*!\[([^\]]*)\]\(([^)]+)\)\s*-->"
)
_INLINE_MATH_WITH_HIDDEN_IMG_RE = re.compile(
    r"(?<!\$)\$(?!\$)((?:\\.|[^$\n])+?)\$(?!\$)\s*"
    r"<!--\s*!\[([^\]]*)\]\(([^)]+)\)\s*-->"
)
_INLINE_MATH_PAREN_WITH_HIDDEN_IMG_RE = re.compile(
    r"\\\((.+?)\\\)\s*"
    r"<!--\s*!\[([^\]]*)\]\(([^)]+)\)\s*-->"
)
_INLINE_MATH_DOLLAR_RE = re.compile(
    r"(?<!\$)\$(?!\$)((?:\\.|[^$\n])+?)\$(?!\$)"
)
_INLINE_MATH_PAREN_RE = re.compile(r"\\\((.+?)\\\)")
_INLINE_TEX_TOKEN_RE = re.compile(
    r"(?<![\w$])"
    r"(?:"
    r"\\[A-Za-z]+(?:_\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}|_[A-Za-z0-9]+|\^\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}|\^[A-Za-z0-9]+)*"
    r"|"
    r"[A-Za-zΑ-Ωα-ω][A-Za-z0-9Α-Ωα-ω]*_(?:\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}|[A-Za-z0-9Α-Ωα-ω]+)"
    r"(?:_\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}|_[A-Za-z0-9]+|\^\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}|\^[A-Za-z0-9Α-Ωα-ω]+)*"
    r")"
    r"(?![\w$])"
)
_BARE_SUBSCRIPT_TOKEN_RE = re.compile(
    r"(?<![A-Za-z0-9$])"
    r"[A-Za-z][A-Za-z0-9]*(?:_[A-Za-z0-9]+)+(?:\^[A-Za-z0-9]+)?"
    r"(?![A-Za-z0-9$])"
)
_SINGLE_LINE_BRACKET_MATH_RE = re.compile(r"^\\\[(.+)\\\]$")
_SINGLE_LINE_DOLLAR_MATH_RE = re.compile(r"^\$\$(.+)\$\$$")
_GREEK_AND_OPERATOR_COMMANDS = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "eta": "η",
    "theta": "θ",
    "lambda": "λ",
    "mu": "μ",
    "rho": "ρ",
    "sigma": "σ",
    "phi": "φ",
    "omega": "ω",
    "Delta": "Δ",
    "Sigma": "Σ",
    "sum": "∑",
    "prod": "∏",
    "cdot": "·",
    "times": "×",
    "le": "≤",
    "leq": "≤",
    "leqslant": "≤",
    "ge": "≥",
    "geq": "≥",
    "geqslant": "≥",
    "neq": "≠",
    "ne": "≠",
    "approx": "≈",
    "infty": "∞",
    "min": "min",
    "max": "max",
    "in": "∈",
    "varepsilon": "ε",
    "epsilon": "ε",
    "pi": "π",
    "Phi": "Φ",
}


def _parse_hidden_image_comment(line: str) -> tuple[str, str] | None:
    m = _HIDDEN_MD_IMAGE_COMMENT_RE.match(line.strip())
    if not m:
        return None
    return m.group(1), m.group(2).strip()


def _try_embed_hidden_comment_line(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    image_max_w_in: float,
    image_max_h_in: float,
) -> bool:
    hidden = _parse_hidden_image_comment(line)
    if not hidden or not base_dir:
        return False
    alt, src = hidden
    if not _resolve_image_path(src, base_dir):
        return False
    _embed_from_image_ref(
        alt,
        src,
        base_dir,
        doc=doc,
        image_max_w_in=image_max_w_in,
        image_max_h_in=image_max_h_in,
    )
    return True


def _image_pixel_size(path: Path) -> tuple[int, int] | None:
    """读取常见位图宽高（像素），失败返回 None。不依赖 Pillow。"""
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if len(raw) >= 24 and raw.startswith(b"\x89PNG\r\n\x1a\n") and raw[12:16] == b"IHDR":
        w = int.from_bytes(raw[16:20], "big")
        h = int.from_bytes(raw[20:24], "big")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 10 and raw[:3] == b"GIF" and raw[3:6] in (b"87a", b"89a"):
        w = int.from_bytes(raw[6:8], "little")
        h = int.from_bytes(raw[8:10], "little")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 4 and raw.startswith(b"\xff\xd8"):
        i = 2
        n = len(raw)
        while i < n:
            if raw[i] != 0xFF:
                i += 1
                continue
            i += 1
            while i < n and raw[i] == 0xFF:
                i += 1
            if i >= n:
                break
            marker = raw[i]
            i += 1
            if marker in (0xD8, 0xD9):
                continue
            if marker == 0xDA:
                break
            if 0xD0 <= marker <= 0xD7:
                continue
            if i + 2 > n:
                break
            seg_len = int.from_bytes(raw[i : i + 2], "big")
            if seg_len < 2:
                break
            i += 2
            if marker in (0xC0, 0xC1, 0xC2) and i + 5 <= n:
                h = int.from_bytes(raw[i + 1 : i + 3], "big")
                w = int.from_bytes(raw[i + 3 : i + 5], "big")
                if w > 0 and h > 0:
                    return w, h
            i += seg_len - 2
    return None


def _fit_image_display_inches(
    px_w: int,
    px_h: int,
    *,
    max_w_in: float,
    max_h_in: float,
) -> tuple[Inches, Inches]:
    """在不超过 max_w / max_h 的前提下等比缩放，使整图落入版面。"""
    if px_w <= 0 or px_h <= 0:
        return Inches(max_w_in), Inches(max_h_in * 0.5)
    aw = max_w_in
    ah = aw * px_h / px_w
    if ah > max_h_in:
        ah = max_h_in
        aw = ah * px_w / px_h
    return Inches(aw), Inches(ah)


def _formula_image_kind(alt: str, src: str) -> str | None:
    """返回 ``block`` / ``inline`` 表示公式图，否则 None（含注释内引用）。"""
    a = alt or ""
    s = src.replace("\\", "/")
    if "math_figures" not in s and "公式" not in a:
        return None
    if "行内" in a:
        return "inline"
    return "block"


def _normalise_math_text_fragment(text: str) -> str:
    out = text
    out = re.sub(r"\\(?:mathrm|text|operatorname)\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", out)
    out = out.replace(r"\left", "").replace(r"\right", "")
    out = out.replace(r"\,", " ").replace(r"\;", " ").replace(r"\:", " ")
    out = out.replace(r"\quad", " ").replace(r"\qquad", " ")
    out = out.replace(r"\_", "_")
    for cmd, value in sorted(_GREEK_AND_OPERATOR_COMMANDS.items(), key=lambda item: len(item[0]), reverse=True):
        out = re.sub(rf"\\{cmd}(?![A-Za-z])", value, out)
    return re.sub(r"\s+", " ", out)


def _extract_formula_tag(text: str) -> tuple[str, str | None]:
    match = re.search(r"\\tag\s*\{([^{}]+)\}", text)
    if not match:
        return text, None
    cleaned = text[: match.start()] + text[match.end() :]
    return cleaned.strip(), match.group(1).strip()


def _strip_math_wrapper(text: str) -> str:
    body = text.strip()
    if body.startswith("$$") and body.endswith("$$") and len(body) >= 4:
        return body[2:-2].strip()
    if body.startswith("$") and body.endswith("$") and len(body) >= 2:
        return body[1:-1].strip()
    if body.startswith(r"\(") and body.endswith(r"\)"):
        return body[2:-2].strip()
    if body.startswith(r"\[") and body.endswith(r"\]"):
        return body[2:-2].strip()
    return body


def _omml_run_xml(text: str) -> str:
    escaped = html.escape(_normalise_math_text_fragment(text))
    return (
        "<m:r>"
        '<m:rPr><m:sty m:val="p"/></m:rPr>'
        f'<m:t xml:space="preserve">{escaped}</m:t>'
        "</m:r>"
    )


def _read_script_arg(text: str, pos: int) -> tuple[str, int]:
    if pos < len(text) and text[pos] == "{":
        depth = 0
        for idx in range(pos, len(text)):
            if text[idx] == "{":
                depth += 1
            elif text[idx] == "}":
                depth -= 1
                if depth == 0:
                    return text[pos + 1 : idx], idx + 1
        return text[pos + 1 :], len(text)
    start = pos
    while pos < len(text) and (text[pos].isalnum() or text[pos] in "\\,"):
        pos += 1
    return text[start:pos], pos


def _normalise_math_atom(text: str) -> str:
    cleaned = _normalise_math_text_fragment(text.strip())
    if cleaned.startswith("\\"):
        key = cleaned[1:]
        return _GREEK_AND_OPERATOR_COMMANDS.get(key, key)
    return cleaned


def _structured_symbol_xml(text: str) -> str:
    first_script = min(
        [idx for idx in (text.find("_"), text.find("^")) if idx != -1],
        default=len(text),
    )
    base = _normalise_math_atom(text[:first_script])
    pos = first_script
    subscript = None
    superscript = None
    while pos < len(text):
        marker = text[pos]
        if marker not in "_^":
            break
        value, pos = _read_script_arg(text, pos + 1)
        if marker == "_":
            subscript = _normalise_math_atom(value)
        else:
            superscript = _normalise_math_atom(value)

    base_xml = _omml_run_xml(base)
    if subscript is not None and superscript is not None:
        return (
            "<m:sSubSup>"
            f"<m:e>{base_xml}</m:e>"
            f"<m:sub>{_omml_run_xml(subscript)}</m:sub>"
            f"<m:sup>{_omml_run_xml(superscript)}</m:sup>"
            "</m:sSubSup>"
        )
    if subscript is not None:
        return f"<m:sSub><m:e>{base_xml}</m:e><m:sub>{_omml_run_xml(subscript)}</m:sub></m:sSub>"
    if superscript is not None:
        return f"<m:sSup><m:e>{base_xml}</m:e><m:sup>{_omml_run_xml(superscript)}</m:sup></m:sSup>"
    return base_xml


def _parse_latex_fragment(text: str) -> str:
    xml, _pos = _parse_latex_sequence(text, 0, stop=None)
    return xml


def _read_balanced_group(text: str, pos: int) -> tuple[str, int]:
    """Read a {...} group starting at pos, returning inner text and next pos."""
    if pos >= len(text) or text[pos] != "{":
        return "", pos
    depth = 0
    start = pos + 1
    for idx in range(pos, len(text)):
        ch = text[idx]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start:idx], idx + 1
    return text[start:], len(text)


def _read_script_value(text: str, pos: int) -> tuple[str, int]:
    """Read one LaTeX atom used as a subscript/superscript value."""
    if pos < len(text) and text[pos] == "{":
        return _read_balanced_group(text, pos)
    if pos < len(text) and text[pos] == "\\":
        start = pos
        pos += 1
        while pos < len(text) and text[pos].isalpha():
            pos += 1
        if pos < len(text) and text[pos] == "{":
            _inner, pos = _read_balanced_group(text, pos)
        return text[start:pos], pos
    start = pos
    while pos < len(text) and (
        text[pos].isalnum() or text[pos] in ",.=+-*/()[]Α-Ωα-ω"
    ):
        pos += 1
    if pos == start and pos < len(text):
        pos += 1
    return text[start:pos], pos


def _wrap_math_script(base_xml: str, sub_xml: str | None, sup_xml: str | None) -> str:
    if sub_xml is not None and sup_xml is not None:
        return (
            "<m:sSubSup>"
            f"<m:e>{base_xml}</m:e>"
            f"<m:sub>{sub_xml}</m:sub>"
            f"<m:sup>{sup_xml}</m:sup>"
            "</m:sSubSup>"
        )
    if sub_xml is not None:
        return f"<m:sSub><m:e>{base_xml}</m:e><m:sub>{sub_xml}</m:sub></m:sSub>"
    if sup_xml is not None:
        return f"<m:sSup><m:e>{base_xml}</m:e><m:sup>{sup_xml}</m:sup></m:sSup>"
    return base_xml


def _read_latex_command_name(text: str, pos: int) -> tuple[str, int]:
    start = pos
    while pos < len(text) and text[pos].isalpha():
        pos += 1
    return text[start:pos], pos


def _parse_latex_command(text: str, pos: int) -> tuple[str, int]:
    """Parse one LaTeX command into an OMML fragment. pos points at backslash."""
    pos += 1
    if pos >= len(text):
        return _omml_run_xml("\\"), pos

    if not text[pos].isalpha():
        ch = text[pos]
        pos += 1
        if ch in {",", ";", ":", " "}:
            return _omml_run_xml(" "), pos
        return _omml_run_xml(ch), pos

    name, pos = _read_latex_command_name(text, pos)

    if name in {"left", "right"}:
        return "", pos

    if name == "frac":
        num, pos = _read_balanced_group(text, pos)
        den, pos = _read_balanced_group(text, pos)
        return (
            "<m:f>"
            f"<m:num>{_parse_latex_fragment(num)}</m:num>"
            f"<m:den>{_parse_latex_fragment(den)}</m:den>"
            "</m:f>"
        ), pos

    if name in {"mathrm", "operatorname", "text"}:
        if pos < len(text) and text[pos] == "{":
            inner, pos = _read_balanced_group(text, pos)
            return _omml_run_xml(inner), pos
        return _omml_run_xml(name), pos

    if name == "mathcal":
        if pos < len(text) and text[pos] == "{":
            inner, pos = _read_balanced_group(text, pos)
            return _omml_run_xml(inner), pos
        return _omml_run_xml(name), pos

    if name in {"quad", "qquad"}:
        return _omml_run_xml(" "), pos

    value = _GREEK_AND_OPERATOR_COMMANDS.get(name)
    if value is not None:
        return _omml_run_xml(value), pos
    return _omml_run_xml(name), pos


def _read_plain_math_text(text: str, pos: int) -> tuple[str, int]:
    start = pos
    while pos < len(text) and text[pos] not in "\\{}_^":
        pos += 1
    return text[start:pos], pos


def _split_plain_script_base(text: str) -> tuple[str, str]:
    """Split a plain run before _/^ so the script binds only to the last symbol."""
    match = re.search(r"([A-Za-z][A-Za-z0-9]*|[Α-Ωα-ω]+)$", text)
    if not match:
        return "", text
    return text[: match.start()], match.group(1)


def _parse_latex_sequence(
    text: str,
    pos: int,
    *,
    stop: str | None,
) -> tuple[str, int]:
    parts: list[str] = []
    while pos < len(text):
        if stop is not None and text[pos] == stop:
            return "".join(parts), pos + 1

        ch = text[pos]
        if ch == "\\":
            atom_xml, pos = _parse_latex_command(text, pos)
        elif ch == "{":
            inner, pos = _read_balanced_group(text, pos)
            atom_xml = _parse_latex_fragment(inner)
        elif ch == "}":
            if stop is None:
                atom_xml = _omml_run_xml(ch)
                pos += 1
            else:
                return "".join(parts), pos + 1
        else:
            plain, pos = _read_plain_math_text(text, pos)
            if plain and pos < len(text) and text[pos] in "_^":
                prefix, base = _split_plain_script_base(plain)
                if prefix:
                    parts.append(_omml_run_xml(prefix))
                atom_xml = _omml_run_xml(base)
            else:
                atom_xml = _omml_run_xml(plain)

        sub_xml = None
        sup_xml = None
        while pos < len(text) and text[pos] in "_^":
            marker = text[pos]
            raw, pos = _read_script_value(text, pos + 1)
            parsed = _parse_latex_fragment(raw)
            if marker == "_":
                sub_xml = parsed
            else:
                sup_xml = parsed
        parts.append(_wrap_math_script(atom_xml, sub_xml, sup_xml))
    return "".join(parts), pos


def _omml_math_xml(text: str) -> str:
    body, _tag = _extract_formula_tag(_strip_math_wrapper(text))
    return (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        + _parse_latex_fragment(body)
        + "</m:oMath>"
    )


def _add_inline_omml_math(paragraph, text: str) -> None:
    paragraph._element.append(parse_xml(_omml_math_xml(text)))


def _make_table_borderless(table) -> None:
    borders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="nil"/>'
        '<w:left w:val="nil"/>'
        '<w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/>'
        '<w:insideH w:val="nil"/>'
        '<w:insideV w:val="nil"/>'
        "</w:tblBorders>"
    )
    table._tbl.tblPr.append(borders)


def _set_cell_width_inches(cell, width_in: float) -> None:
    twips = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            '<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{twips}" w:type="dxa"/>'
        )
    )


def _add_math_block(doc: Document, text: str) -> None:
    body, tag = _extract_formula_tag(_strip_math_wrapper(text))
    if tag:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _make_table_borderless(table)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        _set_cell_width_inches(left, 5.7)
        _set_cell_width_inches(right, 0.7)

        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(6)
        lp.paragraph_format.keep_together = True
        lp._element.append(parse_xml(_omml_math_xml(body)))

        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(3)
        rp.paragraph_format.space_after = Pt(6)
        run = rp.add_run(f"({tag})")
        _set_run_font(run, "宋体", 10.5)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    p._element.append(parse_xml(_omml_math_xml(body)))


def _looks_like_math_text(text: str) -> bool:
    body = _strip_math_wrapper(text)
    if not body:
        return False
    if _INLINE_TEX_TOKEN_RE.search(body):
        return True
    return bool(re.search(r"[A-Za-zΑ-Ωα-ω]", body) and re.search(r"(?:[_^]\{?|[=<>≤≥+\-*/]|\\(?:frac|sum|prod|min|max|cdot|leq|geq))", body))


def _extract_trailing_plain_equation_number(text: str) -> tuple[str, str | None]:
    match = re.search(r"\s+\(([0-9]+[a-z]?)\)\s*$", text.strip())
    if not match:
        return text.strip(), None
    return text.strip()[: match.start()].rstrip(), match.group(1)


def _looks_like_standalone_math_line(text: str) -> bool:
    body, _tag = _extract_trailing_plain_equation_number(text)
    if not body or len(body) < 3:
        return False
    if re.search(r"[\u4e00-\u9fff]", body):
        return False
    if body.startswith(("#", "|", ">", "!", "-", "*", "+", "`")):
        return False
    has_operator = bool(re.search(r"(?:=|<=|>=|\\leq?|\\geq?|[+\-*/])", body))
    has_math_symbol = bool(
        re.search(
            r"(?:[A-Za-z][A-Za-z0-9]*_|\\(?:frac|sum|prod|min|max|cdot|leq|geq|ln|eta|pi|Phi|Delta))",
            body,
        )
    )
    return has_operator and has_math_symbol


def _numbered_math_source(body: str, tag: str | None) -> str:
    stripped = body.strip()
    if not tag:
        return stripped
    return f"{stripped}\\tag{{{tag}}}"


def _parse_formula_numbers_from_manifest(manifest_path: str | Path | None) -> list[str]:
    if manifest_path is None:
        return []
    path = Path(manifest_path)
    if not path.is_file():
        raise FileNotFoundError(path)

    numbers: list[str] = []
    current: dict[str, str] | None = None

    def flush_current() -> None:
        nonlocal current
        if current is None:
            return
        display = current.get("display", "true").strip().lower()
        is_display = display not in {"false", "no", "0"}
        latex = current.get("latex", "").strip()
        number = current.get("number", "").strip()
        formula_id = current.get("id", "").strip()
        if not number and re.fullmatch(r"[0-9]+[a-z]?", formula_id):
            number = formula_id
        if is_display and latex and number:
            numbers.append(number)
        current = None

    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("- "):
            flush_current()
            current = {}
            line = line[2:].strip()
            if not line:
                continue
        if current is None or ":" not in line:
            continue
        key, value = line.split(":", 1)
        key = key.strip()
        if key not in {"id", "number", "latex", "display"}:
            continue
        current[key] = value.strip().strip('"').strip("'")
    flush_current()
    return numbers


def _is_diagram_image(alt: str, src: str) -> bool:
    """mermaid 系统框图 / 流程图等（非公式，用全幅插图尺寸）。"""
    a = alt or ""
    s = src.replace("\\", "/")
    if "mermaid_figures" in s:
        return True
    if a.startswith("图示") or a.startswith("图"):
        return True
    return False


def _span_overlaps(spans: list[tuple[int, int]], start: int, end: int) -> bool:
    return any(not (end <= s or start >= e) for s, e in spans)


def _embed_from_image_ref(
    alt: str,
    src: str,
    base_dir: Path | None,
    *,
    doc: Document | None = None,
    paragraph=None,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> None:
    """按公式 / 框图 / 普通图规则嵌入 PNG（仅公式用小尺寸）。"""
    ipath = _resolve_image_path(src, base_dir) if base_dir else None
    missing = f"[图片缺失: {alt or src}]"
    if not ipath:
        if paragraph is not None:
            paragraph.add_run(missing)
        elif doc is not None:
            doc.add_paragraph().add_run(missing)
        return

    kind = _formula_image_kind(alt, src)
    if kind == "inline":
        p = paragraph
        if p is None and doc is not None:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
        if p is not None:
            _embed_picture_inline(p, ipath, max_h_in=_FORMULA_DISPLAY_MAX_H_IN)
        return

    if doc is None:
        if paragraph is not None:
            paragraph.add_run(missing)
        return

    if kind == "block":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(3)
        _embed_picture_inline(
            p,
            ipath,
            max_h_in=_FORMULA_DISPLAY_MAX_H_IN,
            max_w_in=_FORMULA_BLOCK_MAX_W_IN,
        )
    else:
        _embed_picture(
            doc,
            ipath,
            alt=alt,
            src=src,
            max_w_in=image_max_w_in,
            max_h_in=image_max_h_in,
            center=False,
        )


def _maybe_render_math_md(md_text: str, base_dir: Path) -> str:
    """保留 Markdown 公式源码；写入 Word 时转为可编辑 OMML，而不是 PNG。"""
    return md_text


def _add_math_fallback_block(doc: Document, lines: list[str]) -> None:
    """兼容旧调用名：块级公式写为 Word 原生 OMML。"""
    _add_math_block(doc, "\n".join(ln.rstrip("\n") for ln in lines))


def _embed_picture(
    doc: Document,
    path: Path,
    *,
    alt: str,
    src: str,
    max_w_in: float,
    max_h_in: float,
    center: bool,
) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    try:
        dims = _image_pixel_size(path)
        if dims:
            w_in, h_in = _fit_image_display_inches(
                *dims, max_w_in=max_w_in, max_h_in=max_h_in
            )
            run = p.add_run()
            run.font.bold = False
            run.add_picture(str(path.resolve()), width=w_in, height=h_in)
        else:
            run = p.add_run()
            run.font.bold = False
            run.add_picture(str(path.resolve()), width=Inches(max_w_in))
    except Exception:
        p.add_run(f"[图片无法嵌入: {alt or src} — {path}]")


def _embed_picture_inline(
    paragraph,
    path: Path,
    *,
    max_h_in: float,
    max_w_in: float | None = None,
) -> None:
    try:
        dims = _image_pixel_size(path)
        run = paragraph.add_run()
        run.font.bold = False
        if dims:
            px_w, px_h = dims
            h_in = max_h_in
            w_in = h_in * px_w / px_h if px_h else max_h_in
            if max_w_in is not None and w_in > max_w_in:
                w_in = max_w_in
                h_in = w_in * px_h / px_w if px_w else max_h_in
            run.add_picture(str(path.resolve()), width=Inches(w_in), height=Inches(h_in))
        else:
            run.add_picture(str(path.resolve()), height=Inches(max_h_in))
    except Exception:
        paragraph.add_run(f"[行内公式图缺失: {path}]")


def _add_rich_content_to_paragraph(
    paragraph,
    text: str,
    base_dir: Path | None,
    *,
    formula_inline_max_h_in: float = _FORMULA_DISPLAY_MAX_H_IN,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
    mono: bool = False,
) -> None:
    """同一段内混排文字（**粗体**/`代码`）与公式/插图（含 HTML 注释隐藏引用）。"""
    taken: list[tuple[int, int]] = []
    tokens: list[tuple[int, int, str, tuple]] = []

    for m in _INLINE_MATH_WITH_HIDDEN_IMG_RE.finditer(text):
        tokens.append((m.start(), m.end(), "math_omml", (m.group(1),)))
        taken.append((m.start(), m.end()))

    for m in _INLINE_MATH_PAREN_WITH_HIDDEN_IMG_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "math_omml", (m.group(1),)))
        taken.append((m.start(), m.end()))

    for m in _INLINE_MATH_DOLLAR_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "math_omml", (m.group(1),)))
        taken.append((m.start(), m.end()))

    for m in _INLINE_MATH_PAREN_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "math_omml", (m.group(1),)))
        taken.append((m.start(), m.end()))

    for m in _HIDDEN_MD_IMAGE_COMMENT_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "hidden_img", (m.group(1), m.group(2).strip())))
        taken.append((m.start(), m.end()))

    for m in _MD_IMAGE_RE.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "visible_img", (m.group(1), m.group(2).strip())))
        taken.append((m.start(), m.end()))

    inline_pat = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`)")
    for m in inline_pat.finditer(text):
        if _span_overlaps(taken, m.start(), m.end()):
            continue
        tokens.append((m.start(), m.end(), "inline", (m.group(1),)))
        taken.append((m.start(), m.end()))

    tokens.sort(key=lambda t: t[0])
    pos = 0
    for start, end, kind, payload in tokens:
        if start > pos:
            _add_inline_to_paragraph(paragraph, text[pos:start], mono=mono)
        if kind == "inline":
            token = payload[0]
            if token.startswith("**"):
                run = paragraph.add_run(token[2:-2])
                _set_run_font(run, "宋体", 10.5, bold=True)
            else:
                run = paragraph.add_run(token[1:-1])
                _set_run_font(run, "Consolas", 9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif kind == "math_omml":
            _add_inline_omml_math(paragraph, payload[0])
        else:
            alt, src = payload[0], payload[1]
            _embed_from_image_ref(
                alt,
                src,
                base_dir,
                paragraph=paragraph,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
            )
        pos = end
    if pos < len(text):
        _add_inline_to_paragraph(paragraph, text[pos:], mono=mono)


def _set_run_font(run, name: str = "宋体", size_pt: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _add_plain_run(paragraph, text: str, *, mono: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    _set_run_font(run, "Consolas" if mono else "宋体", 10.5 if not mono else 9)
    if mono:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_text_segment_with_math(paragraph, text: str, *, mono: bool = False) -> None:
    if not text:
        return
    if mono:
        _add_plain_run(paragraph, text, mono=True)
        return

    tokens: list[tuple[int, int, str]] = []
    taken: list[tuple[int, int]] = []
    for pattern in (_INLINE_MATH_DOLLAR_RE, _INLINE_MATH_PAREN_RE):
        for match in pattern.finditer(text):
            if _span_overlaps(taken, match.start(), match.end()):
                continue
            tokens.append((match.start(), match.end(), match.group(1)))
            taken.append((match.start(), match.end()))
    for match in _INLINE_TEX_TOKEN_RE.finditer(text):
        if _span_overlaps(taken, match.start(), match.end()):
            continue
        tokens.append((match.start(), match.end(), match.group(0)))
        taken.append((match.start(), match.end()))
    for match in _BARE_SUBSCRIPT_TOKEN_RE.finditer(text):
        if _span_overlaps(taken, match.start(), match.end()):
            continue
        tokens.append((match.start(), match.end(), match.group(0)))
        taken.append((match.start(), match.end()))

    if not tokens:
        _add_plain_run(paragraph, text)
        return

    tokens.sort(key=lambda item: item[0])
    pos = 0
    for start, end, formula in tokens:
        if start > pos:
            _add_plain_run(paragraph, text[pos:start])
        _add_inline_omml_math(paragraph, formula)
        pos = end
    if pos < len(text):
        _add_plain_run(paragraph, text[pos:])


def _add_inline_to_paragraph(paragraph, text: str, *, mono: bool = False):
    """解析 **粗体**、`行内代码` 与普通文本，写入同一段落。"""
    if not text:
        return
    # 拆分为：粗体、行内代码、普通
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _add_text_segment_with_math(paragraph, text[pos : m.start()], mono=mono)
        token = m.group(1)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, "宋体", 10.5, bold=True)
        else:
            inner = token[1:-1]
            if not mono and _looks_like_math_text(inner):
                _add_inline_omml_math(paragraph, inner)
            else:
                _add_plain_run(paragraph, inner, mono=True)
        pos = m.end()
    if pos < len(text):
        _add_text_segment_with_math(paragraph, text[pos:], mono=mono)


def _add_heading(doc: Document, level: int, text: str):
    """level 1–9 对应 Word 标题 1–标题 9；去除行内标记时保留可读文本。"""
    plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    plain = re.sub(r"`([^`]+)`", r"\1", plain)
    h = doc.add_heading(plain, level=min(max(level, 1), 9))
    for run in h.runs:
        _set_run_font(run, "黑体" if level <= 2 else "宋体")


def _add_body_paragraph(
    doc: Document,
    text: str,
    base_dir: Path | None = None,
    *,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if (
        _MD_IMAGE_RE.search(text)
        or _HIDDEN_MD_IMAGE_COMMENT_RE.search(text)
        or _INLINE_MATH_WITH_HIDDEN_IMG_RE.search(text)
        or _INLINE_MATH_PAREN_WITH_HIDDEN_IMG_RE.search(text)
    ):
        _add_rich_content_to_paragraph(
            p,
            text,
            base_dir,
            image_max_w_in=_DEFAULT_IMAGE_MAX_W_IN,
            image_max_h_in=image_max_h_in,
        )
    else:
        _add_inline_to_paragraph(p, text)
    for run in p.runs:
        if run.font.name in (None, ""):
            _set_run_font(run, "宋体", 10.5)


def _add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    body = "\n".join(lines)
    run = p.add_run(body)
    _set_run_font(run, "Consolas", 9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def _add_list_item(
    doc: Document,
    text: str,
    ordered: bool,
    base_dir: Path | None,
    *,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
    marker: str | None = None,
):
    if ordered:
        # Avoid Word auto-numbering continuing across independent disclosure sections.
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        text = f"{marker or '1.'} {text}"
    else:
        try:
            p = doc.add_paragraph(style="List Bullet")
        except (KeyError, ValueError):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(3)
    if (
        _MD_IMAGE_RE.search(text)
        or _HIDDEN_MD_IMAGE_COMMENT_RE.search(text)
        or _INLINE_MATH_WITH_HIDDEN_IMG_RE.search(text)
        or _INLINE_MATH_PAREN_WITH_HIDDEN_IMG_RE.search(text)
    ):
        _add_rich_content_to_paragraph(
            p,
            text,
            base_dir,
            image_max_w_in=_DEFAULT_IMAGE_MAX_W_IN,
            image_max_h_in=image_max_h_in,
        )
    else:
        _add_inline_to_paragraph(p, text)
    for run in p.runs:
        _set_run_font(run, "宋体", 10.5)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _split_table_cells(line: str) -> list[str]:
    """按列分隔符 ``|`` 拆分表格行，忽略 ``\\(...\\)``、``$...$``、``<!-- -->`` 与 ``\\|`` 内的竖线。"""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]

    cells: list[str] = []
    buf: list[str] = []
    i = 0
    n = len(s)

    while i < n:
        if s.startswith("<!--", i):
            end = s.find("-->", i)
            if end == -1:
                buf.append(s[i:])
                break
            buf.append(s[i : end + 3])
            i = end + 3
            continue

        if s.startswith("\\(", i):
            end = s.find("\\)", i + 2)
            if end == -1:
                buf.append(s[i:])
                break
            buf.append(s[i : end + 2])
            i = end + 2
            continue

        if s[i] == "$":
            if i + 1 < n and s[i + 1] == "$":
                end = s.find("$$", i + 2)
                if end == -1:
                    buf.append(s[i:])
                    break
                buf.append(s[i : end + 2])
                i = end + 2
                continue
            j = i + 1
            while j < n:
                if s[j] == "$" and (j == 0 or s[j - 1] != "\\"):
                    buf.append(s[i : j + 1])
                    i = j + 1
                    break
                j += 1
            else:
                buf.append(s[i:])
                break
            continue

        if s[i] == "\\" and i + 1 < n and s[i + 1] == "|":
            buf.append("\\|")
            i += 2
            continue

        if s[i] == "|":
            cells.append("".join(buf).strip())
            buf = []
            i += 1
            continue

        buf.append(s[i])
        i += 1

    cells.append("".join(buf).strip())
    return cells


def _parse_table_row(line: str) -> list[str]:
    return _split_table_cells(line)


def _is_table_sep(row: list[str]) -> bool:
    if not row:
        return False
    return all(re.match(r"^:?-{3,}:?$", c.strip()) for c in row if c.strip())


def _add_table(doc: Document, rows: list[list[str]], base_dir: Path | None = None):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            if _line_has_embeddable_images(cell_text):
                _add_rich_content_to_paragraph(p, cell_text, base_dir)
            else:
                _add_inline_to_paragraph(p, cell_text)
            for run in p.runs:
                _set_run_font(run, "宋体", 10)


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("─" * 32)
    _set_run_font(run, "宋体", 8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _resolve_image_path(src: str, base_dir: Path | None) -> Path | None:
    if not base_dir:
        return None
    path = (base_dir / src).resolve() if not Path(src).is_absolute() else Path(src)
    return path if path.is_file() else None


def _try_add_image(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> bool:
    m = _MD_IMAGE_RE.match(line.strip())
    if not m or not base_dir:
        return False
    alt, src = m.group(1), m.group(2).strip()
    _embed_from_image_ref(
        alt,
        src,
        base_dir,
        doc=doc,
        image_max_w_in=max_w_in,
        image_max_h_in=max_h_in,
    )
    return True


def _line_has_embeddable_images(line: str) -> bool:
    return bool(
        _MD_IMAGE_RE.search(line)
        or _HIDDEN_MD_IMAGE_COMMENT_RE.search(line)
        or _INLINE_MATH_WITH_HIDDEN_IMG_RE.search(line)
        or _INLINE_MATH_PAREN_WITH_HIDDEN_IMG_RE.search(line)
    )


def _add_paragraph_with_inline_images(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> None:
    """段落内混排文字与公式/插图（含 HTML 注释隐藏引用）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    _add_rich_content_to_paragraph(
        p,
        line,
        base_dir,
        image_max_w_in=max_w_in,
        image_max_h_in=max_h_in,
    )
    for run in p.runs:
        if run.font.name in (None, ""):
            _set_run_font(run, "宋体", 10.5)


def convert_md_to_docx(
    md_text: str,
    base_dir: Path | None,
    *,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
    formula_numbers: list[str] | None = None,
) -> Document:
    doc = Document()
    # 默认正文样式
    try:
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
    except (AttributeError, KeyError):
        pass

    lines = md_text.splitlines()
    i = 0
    para_buf: list[str] = []
    auto_formula_numbers = [n.strip() for n in (formula_numbers or []) if n.strip()]
    formula_number_index = 0

    def flush_paragraph():
        nonlocal para_buf
        if not para_buf:
            return
        # 每行独立成段，避免「（1）…\n（2）…」被空格拼成一段（Word 内不换行）
        for p in para_buf:
            t = p.strip()
            if t:
                _add_body_paragraph(
                    doc,
                    t,
                    base_dir,
                    image_max_h_in=image_max_h_in,
                )
        para_buf = []

    def consume_formula_number(tag: str | None) -> None:
        nonlocal formula_number_index
        if not tag or formula_number_index >= len(auto_formula_numbers):
            return
        remaining = auto_formula_numbers[formula_number_index:]
        if tag in remaining:
            formula_number_index += remaining.index(tag) + 1

    def source_with_auto_formula_number(text: str) -> str:
        nonlocal formula_number_index
        body, tag = _extract_formula_tag(_strip_math_wrapper(text))
        if tag:
            consume_formula_number(tag)
            return text
        if formula_number_index >= len(auto_formula_numbers):
            return text
        tag = auto_formula_numbers[formula_number_index]
        formula_number_index += 1
        return _numbered_math_source(body, tag)

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if line.strip() == "":
            flush_paragraph()
            i += 1
            continue

        # 围栏代码块
        if line.strip().startswith("```"):
            flush_paragraph()
            fence_lang = line.strip()[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            # 定稿 MD 保留 mermaid 源码 + 图示引用：Word 只嵌 PNG，不写源码块。
            # 兼容旧版隐藏 HTML 注释和新版可见 Markdown 图片行。
            if fence_lang.lower() == "mermaid":
                j = i
                while j < len(lines) and lines[j].strip() == "":
                    j += 1
                if j < len(lines):
                    stripped_after_fence = lines[j].strip()
                    cm = _HIDDEN_MD_IMAGE_COMMENT_RE.fullmatch(stripped_after_fence)
                    if not cm:
                        cm = _MD_IMAGE_RE.fullmatch(stripped_after_fence)
                    if cm and _is_diagram_image(cm.group(1), cm.group(2).strip()):
                        continue
            _add_code_block(doc, code_lines)
            continue

        # 块级公式：\[ ... \] + 可选 HTML 注释
        single_bracket_math = _SINGLE_LINE_BRACKET_MATH_RE.match(line.strip())
        if single_bracket_math:
            flush_paragraph()
            _add_math_block(doc, source_with_auto_formula_number(single_bracket_math.group(1).strip()))
            i += 1
            if i < len(lines) and _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip()):
                i += 1
            continue

        if line.strip() == "\\[":
            flush_paragraph()
            i += 1
            math_lines: list[str] = []
            while i < len(lines) and lines[i].strip() != "\\]":
                math_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            if i < len(lines):
                cm = _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip())
                if cm:
                    i += 1
            _add_math_block(doc, source_with_auto_formula_number("\n".join(math_lines)))
            continue

        # 块级公式：$$ ... $$ + 可选 HTML 注释（Word 嵌 PNG；预览见 LaTeX 原文）
        single_dollar_math = _SINGLE_LINE_DOLLAR_MATH_RE.match(line.strip())
        if single_dollar_math:
            flush_paragraph()
            _add_math_block(doc, source_with_auto_formula_number(single_dollar_math.group(1).strip()))
            i += 1
            if i < len(lines) and _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip()):
                i += 1
            continue

        if line.strip() == "$$":
            flush_paragraph()
            i += 1
            math_lines: list[str] = []
            while i < len(lines) and lines[i].strip() != "$$":
                math_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            if i < len(lines):
                cm = _HIDDEN_MD_IMAGE_COMMENT_RE.match(lines[i].strip())
                if cm:
                    i += 1
            _add_math_block(doc, source_with_auto_formula_number("\n".join(math_lines)))
            continue

        # 独立 HTML 注释行（公式图 / mermaid 框图引用）
        if _HIDDEN_MD_IMAGE_COMMENT_RE.fullmatch(line.strip()):
            flush_paragraph()
            _try_embed_hidden_comment_line(
                doc,
                line,
                base_dir,
                image_max_w_in=image_max_w_in,
                image_max_h_in=image_max_h_in,
            )
            i += 1
            continue

        if _looks_like_standalone_math_line(line):
            flush_paragraph()
            body, tag = _extract_trailing_plain_equation_number(line)
            consumed = 1
            if tag is None and i + 1 < len(lines):
                number_line = lines[i + 1].strip()
                number_match = re.fullmatch(r"\(([0-9]+[a-z]?)\)", number_line)
                if number_match:
                    tag = number_match.group(1)
                    consumed = 2
            _add_math_block(doc, source_with_auto_formula_number(_numbered_math_source(body, tag)))
            i += consumed
            continue

        # 图片行或含行内公式/注释的段落
        if _line_has_embeddable_images(line):
            flush_paragraph()
            stripped = line.strip()
            if _MD_IMAGE_RE.fullmatch(stripped) or (
                stripped.startswith("![") and stripped.count("![") == 1
            ):
                _try_add_image(
                    doc,
                    line,
                    base_dir,
                    max_w_in=image_max_w_in,
                    max_h_in=image_max_h_in,
                )
            else:
                _add_paragraph_with_inline_images(
                    doc,
                    line,
                    base_dir,
                    max_w_in=image_max_w_in,
                    max_h_in=image_max_h_in,
                )
            i += 1
            continue

        # 水平线
        if re.match(r"^[\s\-*_]{3,}\s*$", line) and set(line.strip()) <= {"-", "*", "_", " "}:
            flush_paragraph()
            _add_horizontal_rule(doc)
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            flush_paragraph()
            level = len(m.group(1))
            title = m.group(2).strip()
            title = re.sub(r"\s+#+\s*$", "", title)
            _add_heading(doc, level, title)
            i += 1
            continue

        # 引用
        if line.lstrip().startswith("> "):
            flush_paragraph()
            quote = line.lstrip()[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            _add_inline_to_paragraph(p, quote)
            for run in p.runs:
                _set_run_font(run, "宋体", 10.5)
            i += 1
            continue

        # 表格块
        if _is_table_row(line):
            flush_paragraph()
            table_rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                row = _parse_table_row(lines[i])
                if not _is_table_sep(row):
                    table_rows.append(row)
                i += 1
            _add_table(doc, table_rows, base_dir)
            continue

        # 无序列表
        um = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if um:
            flush_paragraph()
            _add_list_item(
                doc,
                um.group(2).strip(),
                ordered=False,
                base_dir=base_dir,
                image_max_h_in=image_max_h_in,
            )
            i += 1
            continue

        # 有序列表
        om = re.match(r"^(\s*)(\d+\.)\s+(.+)$", line)
        if om:
            flush_paragraph()
            _add_list_item(
                doc,
                om.group(3).strip(),
                ordered=True,
                base_dir=base_dir,
                image_max_h_in=image_max_h_in,
                marker=om.group(2),
            )
            i += 1
            continue

        para_buf.append(line)
        i += 1

    flush_paragraph()
    return doc


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description="Markdown → Word（标题样式映射）")
    p.add_argument("-i", "--input", required=True, help="输入 .md 路径")
    p.add_argument("-o", "--output", required=True, help="输出 .docx 路径")
    p.add_argument(
        "--base-dir",
        default=None,
        help="解析 ![](/相对路径) 图片时的根目录（默认使用 .md 所在目录）",
    )
    p.add_argument(
        "--image-max-width-inches",
        type=float,
        default=_DEFAULT_IMAGE_MAX_W_IN,
        metavar="IN",
        help=f"插图最大宽度（英寸，默认 {_DEFAULT_IMAGE_MAX_W_IN}），与高度共同约束等比缩放",
    )
    p.add_argument(
        "--image-max-height-inches",
        type=float,
        default=_DEFAULT_IMAGE_MAX_H_IN,
        metavar="IN",
        help=f"插图最大高度（英寸，默认 {_DEFAULT_IMAGE_MAX_H_IN}），避免竖图仅按宽度缩放后超出单页可视区域",
    )
    p.add_argument(
        "--no-math-render",
        action="store_true",
        help="兼容旧参数；当前默认已不渲染公式图片，公式写为 Word 可编辑 OMML",
    )
    p.add_argument(
        "--math-manifest",
        default=None,
        help="可选公式 manifest（YAML），用于保存后校验公式编号和 OMML 结构",
    )
    p.add_argument(
        "--skip-math-qa",
        action="store_true",
        help="跳过 DOCX 交付 QA（仅用于排障；正式交付不得使用）",
    )
    p.add_argument(
        "--allow-code-style",
        action="store_true",
        help="允许 DOCX 中存在 Consolas/代码样式（仅用于排障；正式交付不得使用）",
    )
    p.add_argument(
        "--min-media-count",
        type=int,
        default=0,
        metavar="N",
        help="要求 DOCX 至少嵌入 N 个 word/media 文件；用于确认 mermaid 图示已进入 Word",
    )
    p.add_argument(
        "--check-formal-text",
        action="store_true",
        help="检查正式交底书正文中是否残留流程说明、交接提示、Agent/脚本名或证据包说明",
    )
    args = p.parse_args(argv)

    in_path = Path(args.input).resolve()
    if not in_path.is_file():
        print(f"错误：找不到输入文件 {in_path}", file=sys.stderr)
        return 1

    base = Path(args.base_dir).resolve() if args.base_dir else in_path.parent
    try:
        md_text = in_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        md_text = in_path.read_text(encoding="utf-8", errors="replace")
        print("警告：输入文件含非 UTF-8 字节，已使用替换字符解码后继续转换。", file=sys.stderr)

    if not args.no_math_render:
        md_text = _maybe_render_math_md(md_text, base)

    try:
        formula_numbers = _parse_formula_numbers_from_manifest(args.math_manifest)
    except Exception as exc:
        print(f"Unable to read math manifest for formula numbering: {exc}", file=sys.stderr)
        return 1

    doc = convert_md_to_docx(
        md_text,
        base_dir=base,
        image_max_w_in=args.image_max_width_inches,
        image_max_h_in=args.image_max_height_inches,
        formula_numbers=formula_numbers,
    )
    out_path = Path(args.output).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已写入: {out_path}")
    if not args.skip_math_qa:
        try:
            from qa_docx_math import check_docx_math, format_report

            report = check_docx_math(
                out_path,
                manifest_path=args.math_manifest,
                allow_code_style=args.allow_code_style,
                min_media_count=max(args.min_media_count, 0),
                check_formal_text=args.check_formal_text,
            )
        except Exception as exc:
            print(f"DOCX 交付 QA 无法执行：{exc}", file=sys.stderr)
            return 1
        print(format_report(report), file=sys.stderr)
        if not report.passed:
            print(
                "错误：DOCX 交付 QA 未通过；请修正 LaTeX/符号表、OMML 或图示生成链路后再交付。",
                file=sys.stderr,
            )
            return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
